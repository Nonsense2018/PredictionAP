#!/usr/bin/env python3
"""Build Appendix — Code.docx with comments and docstrings stripped."""

import ast
import io
import tokenize
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
OUT  = Path.home() / "Downloads" / "Appendix — Code.docx"


# ── Comment + docstring stripper ─────────────────────────────────────────────

def strip(source: str) -> str:
    """Return source with all comments and docstrings removed."""
    try:
        tree = ast.parse(source)
    except SyntaxError:
        return source

    # Collect every line number that belongs to a docstring
    docstring_lines: set[int] = set()
    for node in ast.walk(tree):
        if isinstance(node, (ast.Module, ast.FunctionDef,
                             ast.AsyncFunctionDef, ast.ClassDef)):
            if (node.body
                    and isinstance(node.body[0], ast.Expr)
                    and isinstance(node.body[0].value, ast.Constant)
                    and isinstance(node.body[0].value.value, str)):
                ds = node.body[0]
                for ln in range(ds.lineno, ds.end_lineno + 1):
                    docstring_lines.add(ln)

    # Find inline-comment column positions per line
    comment_col: dict[int, int] = {}
    try:
        toks = list(tokenize.generate_tokens(io.StringIO(source).readline))
        for tok_type, _, tok_start, _, _ in toks:
            if tok_type == tokenize.COMMENT:
                comment_col[tok_start[0]] = tok_start[1]
    except tokenize.TokenError:
        pass

    output: list[str] = []
    for i, line in enumerate(source.splitlines(keepends=True), 1):
        if i in docstring_lines:
            continue
        if i in comment_col:
            col = comment_col[i]
            trimmed = line[:col].rstrip()
            if trimmed:                        # code before the comment
                output.append(trimmed + "\n")
            # else: line was comment-only → drop it
            continue
        output.append(line)

    # Collapse consecutive blank lines — keep at most 1
    cleaned: list[str] = []
    blanks = 0
    for line in output:
        if line.strip() == "":
            blanks += 1
            if blanks <= 1:
                cleaned.append(line)
        else:
            blanks = 0
            cleaned.append(line)

    return "".join(cleaned).strip()


# ── Section definitions ───────────────────────────────────────────────────────

SECTIONS = [
    (
        "Appendix A — Air Quality Data Collection",
        "This code corresponds to the Data Sources and Collection section of the methodology, specifically the EPA AirNow data pipeline.",
        ROOT / "scripts/air/fetch_airnow_history.py",
    ),
    (
        "Appendix B — Meteorological Data Collection",
        "This code corresponds to the Data Sources and Collection section of the methodology, specifically the Open-Meteo weather data pipeline.",
        ROOT / "scripts/met/fetch_openmeteo_history.py",
    ),
    (
        "Appendix C — Wildfire Proximity Data Collection",
        "This code corresponds to the Data Sources and Collection section of the methodology, specifically the NASA EONET wildfire data pipeline.",
        ROOT / "scripts/fire/fetch_eonet_wildfire_data.py",
    ),
    (
        "Appendix D — Feature Engineering",
        "This code corresponds to the Feature Engineering section of the methodology.",
        ROOT / "src/features/build_features.py",
    ),
    (
        "Appendix E — Model Training and Evaluation",
        "This code corresponds to the Model Training and Evaluation section of the methodology.",
        ROOT / "src/models/train_models.py",
    ),
]


# ── Document builder ──────────────────────────────────────────────────────────

def add_page_numbers(doc: Document) -> None:
    footer = doc.sections[0].footer
    para   = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    for tag, attr, val in [
        ("w:fldChar",    "w:fldCharType", "begin"),
        ("w:instrText",  None,            None),
        ("w:fldChar",    "w:fldCharType", "end"),
    ]:
        el = OxmlElement(tag)
        if attr:
            el.set(qn(attr), val)
        else:
            el.text = "PAGE"
        run._r.append(el)


def tight(para) -> None:
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    para.paragraph_format.line_spacing = Pt(12)


def tnr(para, text: str, bold: bool = False, size: int = 12) -> None:
    tight(para)
    run = para.add_run(text)
    run.bold           = bold
    run.font.name      = "Times New Roman"
    run.font.size      = Pt(size)


def main() -> None:
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1)
        sec.right_margin  = Inches(1)

    add_page_numbers(doc)

    for idx, (title, desc, path) in enumerate(SECTIONS):
        if idx > 0:
            doc.add_page_break()

        # Header
        h = doc.add_paragraph()
        tnr(h, title, bold=True)

        # Description
        d = doc.add_paragraph()
        tnr(d, desc)

        # Stripped code
        code = strip(path.read_text(encoding="utf-8"))
        for line in code.splitlines():
            p = doc.add_paragraph()
            tight(p)
            run = p.add_run(line if line.strip() else " ")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        print(f"  {title}: {len(code.splitlines())} lines")

    doc.save(OUT)
    print(f"\nSaved: {OUT}")


if __name__ == "__main__":
    main()
